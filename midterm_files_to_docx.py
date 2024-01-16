from docx import Document
import shutil
import subprocess

# Function to read content from a file
def read_file_content(file_path):
    with open(file_path, 'r') as file:
        content = file.read()
    return content

# Function to create a Word document and paste content
def create_word_document(file_contents, output_filename):
    doc = Document()

    for filename, content in file_contents.items():
        # Add a heading with the filename
        doc.add_heading(filename, level=1)
        # Add the content to the document
        doc.add_paragraph(content)
        # Add a horizontal line between files
        doc.add_paragraph("---------------------------------------------------------------")

    # Save the document
    doc.save(output_filename)

# Specify file paths
header_file_path = 'functions.h'
cpp_file_path = 'functions.cpp'
main_file_path = 'main.cpp'
output_document_path = 'midterm2.docx'
officetopdf_path = "/mnt/e/studentmanual/officetopdf.exe"  # Replace with the actual path

# Read content from files
header_content = read_file_content(header_file_path)
cpp_content = read_file_content(cpp_file_path)
main_content = read_file_content(main_file_path)

# Create a dictionary with filenames and content
file_contents = {
    'functions.h': header_content,
    'functions.cpp': cpp_content,
    'main.cpp': main_content,
}

# Create a Word document and paste content
create_word_document(file_contents, output_document_path)

print(f"Content from {', '.join(file_contents.keys())} copied to {output_document_path}")

# Ask the user whether to convert the Word document to PDF
convert_to_pdf = input("Do you want to convert the Word document to PDF? (yes/no): ").lower()

if convert_to_pdf == 'yes':
    # Specify file paths for conversion
    source_file_path = 'midterm2.docx'
    output_file_path = 'midterm2.pdf'

    # Convert the Word document to PDF using subprocess
    try:
        subprocess.run([officetopdf_path, source_file_path, output_file_path], check=True)
        print(f"Conversion successful. PDF file saved at {output_file_path}")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")
else:
    print("PDF conversion skipped.")